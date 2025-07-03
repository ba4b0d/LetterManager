import os
from tkinter import messagebox, Toplevel, ttk, Label, SUNKEN, W # Importing required Tkinter components
from docx import Document
from docx.shared import Inches # این را اگر لازم داشتید نگه دارید، فعلاً برای جایگزینی مستقیم لازم نیست
import jdatetime # For Persian date conversion
from datetime import datetime # For Gregorian year in filename
import tkinter as tk # For tk.END etc.

# Global references for UI components that helper functions might need to interact with
# These should ideally be passed as arguments or accessed via an App class instance
# For now, we keep them global as they were in your original code, but note this is not ideal.
status_bar = None
progress_window = None
BASE_FONT = ("Arial", 10) # Assuming this is defined elsewhere, or define it here.


def convert_numbers_to_persian(text):
    """Converts English digits in a string to Persian digits."""
    persian_numbers = "۰۱۲۳۴۵۶۷۸۹"
    english_numbers = "0123456789"
    mapping = str.maketrans(english_numbers, persian_numbers)
    return text.translate(mapping)

# CHANGED: More robust text replacement function for DOCX
def replace_text_in_docx(doc_path, replacements):
    """
    Finds and replaces text in a .docx document including headers, footers, and tables.
    Handles placeholders that might be split across multiple runs within a paragraph/cell.
    """
    try:
        document = Document(doc_path)

        def _replace_in_paragraph(paragraph, replacements_dict):
            """Helper to replace text in a single paragraph, handling runs."""
            full_text = paragraph.text
            for old_text, new_text in replacements_dict.items():
                if old_text in full_text:
                    # Simple replacement on paragraph.text if the whole placeholder is there
                    # This might lose formatting across runs if the placeholder was split.
                    # For more complex cases, iterating and rebuilding runs is needed.
                    # For now, let's assume placeholders are mostly contiguous.
                    
                    # A more robust approach for fragmented placeholders:
                    # 1. Collect all runs' text.
                    # 2. Join them to form the full paragraph text.
                    # 3. Perform replacement on the full text.
                    # 4. Clear existing runs.
                    # 5. Add new run(s) with the replaced text, trying to preserve some formatting.
                    
                    # Simplified approach (might lose some formatting for split placeholders but more likely to replace)
                    for run in paragraph.runs:
                        if old_text in run.text: # If a run itself contains the old_text
                            run.text = run.text.replace(old_text, new_text)
                        # More complex: if placeholder spans across runs, we need to handle it differently
                        # For now, we rely on the paragraph.text replacement or partial run replacement.
                        
                    # After checking individual runs, try full paragraph text replacement as a fallback/primary
                    # This often works because paragraph.text concatenates all run texts
                    paragraph.text = paragraph.text.replace(old_text, new_text)


        # Replace in main document body paragraphs
        for paragraph in document.paragraphs:
            _replace_in_paragraph(paragraph, replacements)

        # Replace in tables within the main body
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)

        # Replace in headers and footers
        for section in document.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
                for table in section.header.tables: # Headers can have tables too
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_in_paragraph(paragraph, replacements)
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
                for table in section.footer.tables: # Footers can have tables too
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_in_paragraph(paragraph, replacements)

        document.save(doc_path)
        return True
    except Exception as e:
        messagebox.showerror("خطا در ویرایش Word", f"خطا در جایگزینی متن در سند Word: {e}")
        return False


def sort_column(tree, col, reverse):
    """Sorts a Treeview column."""
    l = [(tree.set(k, col), k) for k in tree.get_children('')]

    if col == "date": # Assuming 'date' is the column with Shamsi dates
        def sort_shamsi_date(item):
            try:
                parts = item[0].split('/')
                # Return a tuple (year, month, day) for correct sorting
                return (int(parts[0]), int(parts[1]), int(parts[2]))
            except (ValueError, IndexError):
                # Handle cases where date format is unexpected, place them at the beginning or end
                return (0, 0, 0) if reverse else (9999, 12, 31)
        l.sort(key=sort_shamsi_date, reverse=reverse)
    elif col == "id": # Treat ID as integer for numeric sort
        l.sort(key=lambda x: int(x[0]) if str(x[0]).isdigit() else (0 if reverse else float('inf')), reverse=reverse)
    else: # Default string sort for others
        l.sort(key=lambda x: x[0].lower() if isinstance(x[0], str) else x[0], reverse=reverse) # Corrected lambda

    # Rearrange items in the Treeview
    for index, (val, k) in enumerate(l):
        tree.move(k, '', index)

# progress_window and show/hide functions (already provided in your helpers.py)
progress_window = None

def show_progress_window(message="در حال پردازش...", root_window_ref=None):
    global progress_window
    if progress_window is None or not progress_window.winfo_exists():
        progress_window = Toplevel(root_window_ref)
        progress_window.title("لطفاً صبر کنید")
        progress_window.geometry("300x100")
        if root_window_ref:
            # Center over the main window
            x = root_window_ref.winfo_x() + (root_window_ref.winfo_width() // 2) - (300 // 2)
            y = root_window_ref.winfo_y() + (root_window_ref.winfo_height() // 2) - (100 // 2)
            progress_window.geometry(f"+{x}+{y}")
        
        progress_window.transient(root_window_ref) # Make it appear on top of root
        progress_window.grab_set() # Disable interaction with other windows
        
        Label(progress_window, text=message, font=BASE_FONT, wraplength=280).pack(pady=20)
        
        # Add a progress bar
        s = ttk.Style()
        s.theme_use('clam') # 'clam' or 'alt' or 'default'
        s.configure("green.Horizontal.TProgressbar", foreground='green', background='green')
        
        progressbar = ttk.Progressbar(progress_window, orient="horizontal", length=200, mode="indeterminate", style="green.Horizontal.TProgressbar")
        progressbar.pack(pady=5)
        progressbar.start(10) # Start animating
        
        progress_window.update_idletasks() # Update GUI immediately

def hide_progress_window():
    global progress_window
    if progress_window is not None and progress_window.winfo_exists():
        progress_window.grab_release()
        progress_window.destroy()
        progress_window = None